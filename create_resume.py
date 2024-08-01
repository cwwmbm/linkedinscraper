import os
from dotenv import load_dotenv
from openai import OpenAI

# Load environment variables from the .env file
load_dotenv()  # Loads the .env file into your environment
# Set the OpenAI API key
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def tailor_resume(job_description, skills_list, resume_text):
    # Prompt template
    prompt_template = """
    **Task:** Adjust the provided resume to perfectly match the given job description by selecting the best four skills for each role from the provided skills list. Ensure the formatting remains consistent and the final document stays within one page.

    ### Instructions:

    1. **Input Files:**
       - **Resume File:** A Word document of the resume.
       - **Skills List File:** A text file containing detailed skills for each role.
       - **Job Description File:** A text file with the job description.

    2. **Steps:**
       - **Read the Job Description:** Extract key skills and requirements from the job description.
       - **Match Skills:** From the provided skills list, select the best four skills for each role that closely align with the job description.
       - **Update the Resume:** Replace the skills in the "Experience" section of the resume with the selected skills. Ensure each role has only four bullet points.
       - **Format the Resume:** Maintain the existing formatting and ensure the final resume fits on one page.

    3. **Formatting Guidelines:**
       - **Margins:** Top, bottom, left, and right margins should be 0.5 inches.
       - **Font:** Use Arial, 10 points for body text, 12 points for section headings, and 14 points for the name.
       - **Line Spacing:** Set line spacing to single (12 points) with no additional space after paragraphs.
       - **Header:** Include the name in bold with a font size of 14 points, the title "Software Engineer" with a font size of 12 points, and contact information aligned to the left.
       - **Sections:** Include sections for Summary, Technical Skills, Experience, and Education. Each section should have a heading in bold with a font size of 12 points.
       - **Experience Section:** Ensure each role has four bullet points with selected skills. Dates should be aligned to the right of job titles.

    4. **Output:** A formatted Word document with the tailored resume.

    ### Example Implementation in Python (for formatting reference):

    ```python
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Create a new Document
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)   # 0.5 inches
        section.bottom_margin = Pt(36)  # 0.5 inches
        section.left_margin = Pt(36)  # 0.5 inches
        section.right_margin = Pt(36)  # 0.5 inches

    # Set default font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(12)

    # Helper function to add headings with formatting
    def add_heading(text, level, bold=False, font_size=12):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return heading

    # Add header with name and title
    add_heading('Chris Phillips', 0, bold=True, font_size=14)
    add_heading('Software Engineer', 1, font_size=12)

    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.add_run('678-409-8713 | phillipsachris@gmail.com | https://cphillips.dev\n')
    contact_info.add_run('https://www.linkedin.com/in/chris-a-phillips | https://github.com/chris-a-phillips')
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Summary
    add_heading('Summary', level=1, bold=True, font_size=12)
    summary = (
        "Results-driven Software Engineer with over three years of experience in developing scalable and "
        "efficient web applications. Adept at solving complex problems and delivering high-quality software "
        "solutions. Proven track record of working collaboratively in Agile and fast-paced environments to "
        "drive innovation and achieve project goals."
    )
    doc.add_paragraph(summary)

    # Add Technical Skills
    add_heading('Technical Skills', level=1, bold=True, font_size=12)
    skills = (
        "Languages: JavaScript, Python, TypeScript, Java, SQL, Bash, HTML/CSS\n"
        "Frameworks/Runtimes: Node, React, Next.js, Express, Mongoose, Luigi, Django\n"
        "Dev Tools: Git, AWS, Docker, Postman, Jenkins\n"
        "Databases: MySQL, MongoDB, PostgreSQL\n"
        "Tools: Git, Docker, GitHub Actions, Jenkins, Jira, Confluence, Kubernetes, Terraform, Datadog"
    )
    doc.add_paragraph(skills)

    # Add Experience
    add_heading('Experience', level=1, bold=True, font_size=12)

    # Example role update for Built Technologies
    built_tech = add_heading('Software Engineer | Built Technologies', level=2, bold=True, font_size=12)
    built_tech.add_run('\tFebruary 2022 – Present')
    doc.add_paragraph(
        "- Selected Skill 1\n"
        "- Selected Skill 2\n"
        "- Selected Skill 3\n"
        "- Selected Skill 4"
    )

    # Repeat similar structure for other roles

    # Add Education
    add_heading('Education', level=1, bold=True, font_size=12)

    # General Assembly
    ga_edu = add_heading('General Assembly | Remote', level=2, bold=True, font_size=12)
    ga_edu.add_run('\tOctober 2020 - December 2020')
    doc.add_paragraph('Software Engineering Immersive')

    # University of North Carolina at Greensboro
    uncg = add_heading('University of North Carolina at Greensboro | Greensboro, North Carolina', level=2, bold=True, font_size=12)
    uncg.add_run('\tAugust 2013 - May 2016')
    doc.add_paragraph('Psychology')

    # Save the document
    file_path = "Chris_Phillips_Resume_Tailored.docx"
    doc.save(file_path)
    ```

    ### Skills List:
    {skills_list}

    ### Job Description:
    {job_description}
    """

    # Fill in the prompt with job description and skills list
    prompt = prompt_template.format(skills_list=skills_list, job_description=job_description)

    # Send the request to ChatGPT
    response = client.chat.completions.create(model="gpt-4",
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt},
    ])

    # Extract the tailored resume content
    tailored_resume_content = response.choices[0].message.content

    return tailored_resume_content

def save_resume_to_docx(resume_text, output_path):
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)   # 0.5 inches
        section.bottom_margin = Pt(36)  # 0.5 inches
        section.left_margin = Pt(36)  # 0.5 inches
        section.right_margin = Pt(36)  # 0.5 inches

    # Set default font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(12)

    # Split resume text into lines and add them to the document
    lines = resume_text.split('\n')
    for line in lines:
        if line.strip() != '':
            doc.add_paragraph(line)

    # Save the document
    doc.save(output_path)

# Example usage
with open('./data/input/job_description.txt', 'r') as file:
    job_description = file.read()

with open('./data/input/skills_list.txt', 'r') as file:
    skills_list = file.read()

with open('./data/input/resume.docx', 'rb') as file:
    resume_text = file.read()

tailored_resume = tailor_resume(job_description, skills_list, resume_text)
output_path = './data/output/Chris Phillips Resume.docx'
save_resume_to_docx(tailored_resume, output_path)

print(f"Tailored resume saved to {output_path}")






