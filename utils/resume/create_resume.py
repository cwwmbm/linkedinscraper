from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement


def set_font(run, name, size, color=None, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p_element = p._element
    p_pr = p_element.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'auto')
    p_borders.append(bottom_border)
    p_pr.append(p_borders)

def add_paragraph_with_format(doc, text, style=None, alignment=None, space_before=None, space_after=1, left_indent=None, first_line_indent=None, line_spacing=None):
    p = doc.add_paragraph(text, style)
    if alignment:
        p.alignment = alignment
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)  # Always set space_after to a smaller value unless specified otherwise
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    return p


def create_general_resume(contact, summary, technical_skills, experience, education):
    # Create a new document
    doc = Document()

    # Title and Contact Information
    p = add_paragraph_with_format(doc, contact['name'], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 13.5, bold=True)
    add_paragraph_with_format(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p = add_paragraph_with_format(doc, contact['title'], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10)

    p = add_paragraph_with_format(doc, f"{contact['number']} | {contact['email']} | {contact['portfolio_url']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 8.5)

    p = add_paragraph_with_format(doc, f"{contact['linkedin_url']}  |  {contact['github_url']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 8.5)

    add_divider(doc)

    # Summary
    p = add_paragraph_with_format(doc, "Summary", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True)
    p.paragraph_format.space_after = Pt(0)

    p = add_paragraph_with_format(doc, summary, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    set_font(p.runs[0], 'Arial', 8.5)
    p.paragraph_format.space_before = Pt(0)

    add_divider(doc)

    # Technical Skills
    for skill, details in technical_skills.items():
        p = add_paragraph_with_format(doc, skill, style='Normal', space_after=2)
        set_font(p.runs[0], 'Arial', 8.5, bold=True)
        run = p.add_run(details)
        set_font(run, 'Arial', 8.5)


    add_divider(doc)

    # # # EXPERIENCE
    # Experience Header
    p = add_paragraph_with_format(doc, "Experience", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True)
    p = add_paragraph_with_format(doc, "", line_spacing=WD_LINE_SPACING.SINGLE)
    p.paragraph_format.space_after = Pt(0)

    # Experience Content
    for i in range(len(experience)):
        # Company and Dates
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        run_bold = p.add_run(experience[i]['company'])
        set_font(run_bold, 'Arial', 8.5, bold=True)
        run_italic = p.add_run(experience[i]['dates'])
        set_font(run_italic, 'Arial', 8.5, italic=True)
        p.paragraph_format.space_after = Pt(0)

        # Points
        for j in range(len(experience[i]['points'])):
            text = experience[i]['points'][j]
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.2)  # Adjust this value to control the space between the bullet and the text
            p.paragraph_format.first_line_indent = Inches(-0.15)  # Adjust this value to control the space between the bullet and the text
            set_font(p.runs[0], 'Arial', 8.5)
            p.paragraph_format.space_after = Pt(0)

        # Add a blank line for separation
        if i < (len(experience) - 1):
            doc.add_paragraph("")



    add_divider(doc)

    # Education
    p = add_paragraph_with_format(doc, "Education", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True, color=RGBColor(0x1B, 0x1F, 0x22))

    for edu in education:
        # Institution and Dates
        p = add_paragraph_with_format(doc, edu['institution'] + edu['dates'], line_spacing=WD_LINE_SPACING.SINGLE)
        set_font(p.runs[0], 'Arial', 8.5, bold=True)
        set_font(p.runs[-1], 'Arial', 8.5, italic=True, color=RGBColor(0x1B, 0x1F, 0x22))

        # Degree
        p = add_paragraph_with_format(doc, edu['degree'], style='Normal')
        set_font(p.runs[0], 'Arial', 8.5)

        doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)


    # Set document section details
    section = doc.sections[0]
    section.start_type = 2  # New page
    section.orientation = 0  # Portrait
    section.page_width = 7772400
    section.page_height = 10058400
    section.left_margin = 914400
    section.right_margin = 914400
    section.top_margin = 457200  # 0.5 inches
    section.bottom_margin = 457200  # 0.5 inches
    section.header_distance = 457200
    section.footer_distance = 457200
    section.gutter = 0


    # Save document
    doc.save("./data/output/Chris Phillips Resume.docx")

    # job skills max characters per bullet point is 110
    # summary max lines is 5 and max characters per line is 105
    # technical skills max characters is 121


def create_gpt_resume(contact, summary, technical_skills, experience, education, path):
    # Create a new document
    doc = Document()

    # Title and Contact Information
    p = add_paragraph_with_format(doc, contact['name'], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 13.5, bold=True)
    add_paragraph_with_format(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p = add_paragraph_with_format(doc, contact['title'], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10)

    p = add_paragraph_with_format(doc, f"{contact['number']} | {contact['email']} | {contact['portfolio_url']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 8.5)

    p = add_paragraph_with_format(doc, f"{contact['linkedin_url']}  |  {contact['github_url']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 8.5)

    add_divider(doc)

    # Summary
    p = add_paragraph_with_format(doc, "Summary", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True)
    p.paragraph_format.space_after = Pt(0)

    p = add_paragraph_with_format(doc, summary, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    set_font(p.runs[0], 'Arial', 8.5)
    p.paragraph_format.space_before = Pt(0)

    add_divider(doc)

    # Technical Skills
    for skill, details in technical_skills.items():
        p = add_paragraph_with_format(doc, skill, style='Normal', space_after=2)
        set_font(p.runs[0], 'Arial', 8.5, bold=True)
        run = p.add_run(details)
        set_font(run, 'Arial', 8.5)


    add_divider(doc)

    # # # EXPERIENCE
    # Experience Header
    p = add_paragraph_with_format(doc, "Experience", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True)
    p = add_paragraph_with_format(doc, "", line_spacing=WD_LINE_SPACING.SINGLE)
    p.paragraph_format.space_after = Pt(0)

    # Experience Content
    for i in range(len(experience)):
        # Company and Dates
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        run_bold = p.add_run(experience[i]['company'])
        set_font(run_bold, 'Arial', 8.5, bold=True)
        run_italic = p.add_run(experience[i]['dates'])
        set_font(run_italic, 'Arial', 8.5, italic=True)
        p.paragraph_format.space_after = Pt(0)

        # Points
        for j in range(len(experience[i]['points'])):
            text = experience[i]['points'][j]
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.2)  # Adjust this value to control the space between the bullet and the text
            p.paragraph_format.first_line_indent = Inches(-0.15)  # Adjust this value to control the space between the bullet and the text
            set_font(p.runs[0], 'Arial', 8.5)
            p.paragraph_format.space_after = Pt(0)

        # Add a blank line for separation
        if i < (len(experience) - 1):
            doc.add_paragraph("")



    add_divider(doc)

    # Education
    p = add_paragraph_with_format(doc, "Education", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.runs[0], 'Arial', 10.5, bold=True, color=RGBColor(0x1B, 0x1F, 0x22))

    for edu in education:
        # Institution and Dates
        p = add_paragraph_with_format(doc, edu['institution'] + edu['dates'], line_spacing=WD_LINE_SPACING.SINGLE)
        set_font(p.runs[0], 'Arial', 8.5, bold=True)
        set_font(p.runs[-1], 'Arial', 8.5, italic=True, color=RGBColor(0x1B, 0x1F, 0x22))

        # Degree
        p = add_paragraph_with_format(doc, edu['degree'], style='Normal')
        set_font(p.runs[0], 'Arial', 8.5)

        doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)


    # Set document section details
    section = doc.sections[0]
    section.start_type = 2  # New page
    section.orientation = 0  # Portrait
    section.page_width = 7772400
    section.page_height = 10058400
    section.left_margin = 914400
    section.right_margin = 914400
    section.top_margin = 457200  # 0.5 inches
    section.bottom_margin = 457200  # 0.5 inches
    section.header_distance = 457200
    section.footer_distance = 457200
    section.gutter = 0


    # Save document
    # doc.save("./data/output/Chris Phillips Resume.docx")
    doc.save(path)

    # job skills max characters per bullet point is 110
    # summary max lines is 5 and max characters per line is 105
    # technical skills max characters is 121